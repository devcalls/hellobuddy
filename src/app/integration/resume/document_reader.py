from pathlib import Path


class UnsupportedDocumentError(Exception):
    pass


class DocumentReadError(Exception):
    pass


class DocumentReader:

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".txt",
    }

    def read(
        self,
        file_path: str,
    ) -> str:

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"Resume file not found: {file_path}")

        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise UnsupportedDocumentError(
                f"Unsupported resume format: {extension}. "
                f"Supported formats: "
                f"{', '.join(sorted(self.SUPPORTED_EXTENSIONS))}"
            )

        if extension == ".pdf":
            return self._read_pdf(path)

        if extension == ".docx":
            return self._read_docx(path)

        return self._read_txt(path)

    def _read_pdf(
        self,
        path: Path,
    ) -> str:

        try:
            import fitz
        except ImportError as exc:
            raise DocumentReadError(
                "PyMuPDF is required for PDF parsing. "
                "Install with: pip install pymupdf"
            ) from exc

        try:
            document = fitz.open(path)

            pages = []

            for page in document:
                text = page.get_text("text")

                if text:
                    pages.append(text.strip())

            document.close()

            return "\n\n".join(pages).strip()

        except Exception as exc:
            raise DocumentReadError(f"Unable to read PDF: {path}") from exc

    def _read_docx(
        self,
        path: Path,
    ) -> str:

        try:
            from docx import Document
        except ImportError as exc:
            raise DocumentReadError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx"
            ) from exc

        try:
            document = Document(path)

            blocks = []

            # Paragraphs
            for paragraph in document.paragraphs:

                text = paragraph.text.strip()

                if text:
                    blocks.append(text)

            # Tables
            #
            # We don't use tables as resume sections yet,
            # but extracting their text prevents losing
            # content from resumes that use tables.
            for table in document.tables:

                for row in table.rows:

                    row_text = " | ".join(
                        cell.text.strip() for cell in row.cells if cell.text.strip()
                    )

                    if row_text:
                        blocks.append(row_text)

            return "\n".join(blocks).strip()

        except Exception as exc:
            raise DocumentReadError(f"Unable to read DOCX: {path}") from exc

    def _read_txt(
        self,
        path: Path,
    ) -> str:

        try:
            return path.read_text(encoding="utf-8").strip()

        except UnicodeDecodeError as exc:
            raise DocumentReadError(
                f"Unable to decode text file as UTF-8: {path}"
            ) from exc
